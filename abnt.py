"""
Streamlit ABNT Formatter (DOCX → DOCX)
--------------------------------------
- Aplica formatação ABNT em .docx e permite baixar o resultado.
- Foco: **NBR 14724** (margens, fonte, espaçamento, recuo, alinhamento, numeração) + **tabelas/figuras/citações longas** e **Referências (NBR 6023) – modo assistido**.

Como rodar (local):
1) pip install -r requirements.txt
2) streamlit run app.py

requirements.txt:
------------------
streamlit
python-docx
lxml

Notas:
- O app não exporta PDF (faça no Word/LibreOffice após download).
- "Citação longa" e "Referências NBR 6023" têm muita nuance: oferecemos **marcação assistida** e **templates**.
"""

import io
from typing import Optional, List, Tuple
import re
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================
# ABNT Helper Functions
# =====================

def set_page_margins(doc: Document, top_cm=3.0, left_cm=3.0, right_cm=2.0, bottom_cm=2.0):
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.bottom_margin = Cm(bottom_cm)


def configure_default_style(doc: Document, font_name="Times New Roman", font_size_pt=12, line_spacing=1.5,
                            first_line_indent_cm=1.25, justify=True):
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    pf = normal.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(first_line_indent_cm)


def style_all_paragraphs(doc: Document, justify=True, first_line_indent_cm=1.25):
    for p in doc.paragraphs:
        text = p.text.strip()
        # Headings: alinhar à esquerda, sem recuo
        if p.style and p.style.name.lower().startswith("heading"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            if justify:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def uppercase_heading_runs(paragraph):
    for run in paragraph.runs:
        run.text = run.text.upper()


def configure_heading_styles(doc: Document, h1_caps=True, h2_caps=True, h3_caps=True):
    heading_map = {"Heading 1": h1_caps, "Heading 2": h2_caps, "Heading 3": h3_caps}
    for p in doc.paragraphs:
        name = p.style.name if p.style is not None else ""
        if name in heading_map and heading_map[name]:
            uppercase_heading_runs(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)


def add_page_number_to_footer(doc: Document, position="right"):
    for section in doc.sections:
        footer = section.footer
        para = footer.add_paragraph() if len(footer.paragraphs) == 0 else footer.paragraphs[0]
        para.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT
        }.get(position, WD_ALIGN_PARAGRAPH.RIGHT)

        run = para.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = ' PAGE '
        fld_separate = OxmlElement('w:fldChar'); fld_separate.set(qn('w:fldCharType'), 'separate')
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr_text); run._r.append(fld_separate); run._r.append(fld_end)


def remove_extra_blank_lines(doc: Document):
    i = 0
    while i < len(doc.paragraphs) - 1:
        if not doc.paragraphs[i].text.strip() and not doc.paragraphs[i+1].text.strip():
            try:
                p = doc.paragraphs[i+1]._element
                p.getparent().remove(p)
                continue
            except Exception:
                pass
        i += 1

# ----------
# Tabelas
# ----------

def _set_row_cant_split(row):
    # <w:trPr><w:cantSplit/></w:trPr>
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)


def prevent_table_row_split_and_repeat_header(table):
    # Impede quebra de linha da linha na página seguinte e repete cabeçalho
    for i, row in enumerate(table.rows):
        _set_row_cant_split(row)
        if i == 0:
            trPr = row._tr.get_or_add_trPr()
            hdr = OxmlElement('w:tblHeader')
            trPr.append(hdr)
    # Melhor legibilidade de largura (evita autofit extremo)
    try:
        tblPr = table._tbl.get_or_add_tblPr()
        tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
    except Exception:
        pass


def center_paragraphs_with_drawings(doc: Document):
    # Centraliza parágrafos que contêm imagens (desenhos)
    for p in doc.paragraphs:
        if p._element.xpath('.//w:drawing'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)

# ----------
# Citações Longas (NBR 10520 – assistida)
# ----------
# Marcadores no texto: [[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]
# O app converterá os parágrafos marcados em bloco com: recuo 4 cm, fonte 10 pt, espaçamento simples, sem aspas.


def apply_long_quote_style(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(4)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.style.font.size = Pt(10) if paragraph.style and paragraph.style.font else None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def process_long_quote_markers(doc: Document) -> int:
    in_block = False
    changed = 0
    for p in doc.paragraphs:
        txt = p.text
        if '[[CITACAO_LONGA]]' in txt:
            in_block = True
            p.text = txt.replace('[[CITACAO_LONGA]]', '')
            apply_long_quote_style(p)
            changed += 1
            continue
        if '[[/CITACAO_LONGA]]' in txt:
            in_block = False
            p.text = txt.replace('[[/CITACAO_LONGA]]', '')
            apply_long_quote_style(p)
            changed += 1
            continue
        if in_block:
            apply_long_quote_style(p)
            changed += 1
    return changed

# ----------
# Referências (NBR 6023 – assistidas)
# ----------
# Opção 1: marcar bloco de referências com [[REFERENCIAS]] ... [[/REFERENCIAS]] para aplicar recuo francês
# Opção 2: utilizar gerador de referência por tipo (Livro, Artigo, Site)


def apply_references_block_format(doc: Document, first_line_hanging_cm=1.25, line_spacing=1.0, space_between_pts=6):
    in_refs = False
    count = 0
    for p in doc.paragraphs:
        t = p.text
        if '[[REFERENCIAS]]' in t:
            in_refs = True
            p.text = t.replace('[[REFERENCIAS]]', '')
        elif '[[/REFERENCIAS]]' in t:
            in_refs = False
            p.text = t.replace('[[/REFERENCIAS]]', '')
        if in_refs or '[[/REFERENCIAS]]' in t:
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(first_line_hanging_cm)
            pf.line_spacing = line_spacing
            pf.space_after = Pt(space_between_pts)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            count += 1
    return count


def format_reference_livro(autor_sobrenome: str, autor_iniciais: str, titulo: str, ed: Optional[str], local: str, editora: str, ano: str):
    # SOBRENOME, Iniciais. Título: subtítulo. ed. Local: Editora, ano.
    ed_str = f" {ed}." if ed else "."
    return f"{autor_sobrenome.upper()}, {autor_iniciais}. {titulo}.{ed_str} {local}: {editora}, {ano}."


def format_reference_artigo(autor_sobrenome: str, autor_iniciais: str, titulo: str, periodico: str, volume: str, numero: Optional[str], paginas: str, ano: str):
    num = f"({numero})" if numero else ""
    return f"{autor_sobrenome.upper()}, {autor_iniciais}. {titulo}. {periodico}, v. {volume} {num}, p. {paginas}, {ano}."


def format_reference_site(autor_sobrenome: Optional[str], autor_iniciais: Optional[str], titulo: str, site: str, url: str, acesso_data: str, ano: Optional[str]=None):
    autor = f"{autor_sobrenome.upper()}, {autor_iniciais}. " if (autor_sobrenome and autor_iniciais) else ""
    ano_str = f" {ano}." if ano else "."
    return f"{autor}{titulo}. {site}. Disponível em: <{url}>. Acesso em: {acesso_data}.{ano_str}"

# ----------
# Figuras (centralizar e legenda abaixo) e Tabelas (título acima, fonte abaixo)
# ----------


def add_caption_after_paragraph(p, text: str, italic=False):
    cap = p.insert_paragraph_after(text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    if italic and cap.runs:
        cap.runs[0].italic = True
    return cap


def ensure_captions(doc: Document, add_fig_captions: bool, add_tab_captions: bool):
    fig_n = 0
    tab_n = 0
    for block in doc.element.body:
        tag = block.tag
        if tag.endswith('}p'):
            p = next((pp for pp in doc.paragraphs if pp._p is block), None)
            if p is not None and p._element.xpath('.//w:drawing'):
                # imagem encontrada
                if add_fig_captions:
                    fig_n += 1
                    add_caption_after_paragraph(p, f"Figura {fig_n} – Descrição da figura", italic=False)
        elif tag.endswith('}tbl'):
            # tabela encontrada
            tbl = next((t for t in doc.tables if t._tbl is block), None)
            if tbl is not None:
                prevent_table_row_split_and_repeat_header(tbl)
                if add_tab_captions:
                    # Título acima e fonte abaixo (placeholder)
                    # Inserir título acima
                    first_p = doc.paragraphs[0]
                    new_p_above = first_p._element.__class__('w:p')
                    block.addprevious(new_p_above)
                    p_obj_above = None
                    for pp in doc.paragraphs:
                        if pp._p is new_p_above:
                            p_obj_above = pp; break
                    if p_obj_above:
                        p_obj_above.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_obj_above.paragraph_format.first_line_indent = Cm(0)
                        p_obj_above.add_run(f"Tabela {tab_n+1} – Título da tabela")
                    # Fonte abaixo
                    new_p_below = first_p._element.__class__('w:p')
                    block.addnext(new_p_below)
                    p_obj_below = None
                    for pp in doc.paragraphs:
                        if pp._p is new_p_below:
                            p_obj_below = pp; break
                    if p_obj_below:
                        p_obj_below.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_obj_below.paragraph_format.first_line_indent = Cm(0)
                        p_obj_below.add_run("Fonte: elaboração própria.")
                    tab_n += 1

# =====================
# Pipeline principal
# =====================

def apply_abnt_formatting(doc: Document,
                          h1_caps=True,
                          h2_caps=True,
                          h3_caps=True,
                          justify=True,
                          footer_page_numbers=True,
                          first_line_indent_cm=1.25,
                          center_images=True,
                          auto_captions_fig=True,
                          auto_captions_tab=True,
                          format_refs_block=True):
    set_page_margins(doc)
    configure_default_style(doc, line_spacing=1.5, first_line_indent_cm=first_line_indent_cm)
    style_all_paragraphs(doc, justify=justify, first_line_indent_cm=first_line_indent_cm)
    configure_heading_styles(doc, h1_caps=h1_caps, h2_caps=h2_caps, h3_caps=h3_caps)

    if center_images:
        center_paragraphs_with_drawings(doc)

    # Tabelas: impedir quebra de linha e repetir cabeçalho; adicionar legendas (opcional)
    for t in doc.tables:
        prevent_table_row_split_and_repeat_header(t)

    ensure_captions(doc, add_fig_captions=auto_captions_fig, add_tab_captions=auto_captions_tab)

    # Citações longas via marcadores
    process_long_quote_markers(doc)

    # Bloco de referências formatado
    if format_refs_block:
        apply_references_block_format(doc)

    if footer_page_numbers:
        add_page_number_to_footer(doc)

    remove_extra_blank_lines(doc)
    return doc

# =====================
# Streamlit UI
# =====================

st.set_page_config(page_title="ABNT Formatter (DOCX)", page_icon="📄", layout="centered")

st.title("📄 Automatizador ABNT para DOCX – Completo")
st.caption("Margens, fonte, parágrafos, títulos, numeração, **tabelas**, **figuras**, **citações longas** e **Referências (assistido)**.")

with st.expander("⚙️ Opções de formatação", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        do_h1 = st.checkbox("TÍTULOS NÍVEL 1 EM CAIXA ALTA", value=True)
        do_h2 = st.checkbox("TÍTULOS NÍVEL 2 EM CAIXA ALTA", value=True)
        do_h3 = st.checkbox("TÍTULOS NÍVEL 3 EM CAIXA ALTA", value=True)
        do_justify = st.checkbox("Justificar parágrafos", value=True)
        first_indent = st.number_input("Recuo da 1ª linha (cm)", min_value=0.0, max_value=3.0, value=1.25, step=0.25)
    with col2:
        do_pagenum = st.checkbox("Inserir número de página no rodapé (direita)", value=True)
        center_imgs = st.checkbox("Centralizar imagens (parágrafos com figuras)", value=True)
        auto_fig = st.checkbox("Adicionar legenda abaixo das figuras (placeholder)", value=False)
        auto_tab = st.checkbox("Adicionar título acima e fonte abaixo das tabelas (placeholder)", value=False)
        refs_block = st.checkbox("Formatar bloco de referências (recuo francês)", value=True)

st.markdown(
    """**Marcadores úteis no texto do DOCX** (opcionais):
- `[[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]` → aplica NBR 10520: recuo 4 cm, fonte ~10 pt, espaçamento simples, sem aspas.
- `[[REFERENCIAS]] ... [[/REFERENCIAS]]` → aplica recuo francês, espaçamento simples e espaço entre entradas.
    """
)

uploaded = st.file_uploader("Envie seu arquivo .docx", type=["docx"]) 

if uploaded is not None:
    try:
        doc = Document(uploaded)
    except Exception as e:
        st.error(f"Erro ao abrir DOCX: {e}")
        st.stop()

    st.info("Clique em **Aplicar ABNT** para processar. Você poderá baixar o DOCX formatado.")

    if st.button("✨ Aplicar ABNT"):
        try:
            formatted = apply_abnt_formatting(
                doc,
                h1_caps=do_h1,
                h2_caps=do_h2,
                h3_caps=do_h3,
                justify=do_justify,
                footer_page_numbers=do_pagenum,
                first_line_indent_cm=float(first_indent),
                center_images=center_imgs,
                auto_captions_fig=auto_fig,
                auto_captions_tab=auto_tab,
                format_refs_block=refs_block,
            )
            out = io.BytesIO(); formatted.save(out); out.seek(0)
            base_name = uploaded.name.replace('.docx', '').replace('.DOCX', '')
            st.success("Arquivo formatado com sucesso!")
            st.download_button(
                label="⬇️ Baixar DOCX formatado",
                data=out,
                file_name=f"{base_name}_ABNT.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.caption("Dica: revise capa/folha de rosto, listas pré-textuais e sumário no Word/LibreOffice e exporte para PDF.")
        except Exception as e:
            st.error(f"Falha ao formatar: {e}")

st.divider()

st.subheader("Cobertura desta automação")
st.markdown(
    """
**Layout geral (NBR 14724):**
- **Margens**: 3 cm (superior/esquerda) e 2 cm (direita/inferior)
- **Fonte**: Times New Roman, 12 pt (texto)
- **Parágrafos**: 1,5; sem espaços extras; **recuo 1,25 cm**; justificado (opcional)
- **Títulos (Heading 1/2/3)**: caixa alta opcional; esquerda; sem recuo
- **Páginas**: numeração no rodapé (direita) com campo automático `PAGE`

**Tabelas e Figuras:**
- **Tabelas**: impede que a **linha quebre entre páginas** e repete **cabeçalho**; opção de **Título acima** e **Fonte abaixo** (placeholders).
- **Figuras**: centraliza parágrafo com imagem; opção de **Legenda abaixo** numerada (placeholder).

**Citações (NBR 10520):**
- **Citação curta**: siga no corpo entre aspas (até ~3 linhas) – revisão manual.
- **Citação longa**: use os marcadores `[[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]` para aplicar bloco ABNT.

**Referências (NBR 6023):**
- Use `[[REFERENCIAS]] ... [[/REFERENCIAS]]` para aplicar **recuo francês**, espaçamento simples e espaço entre entradas.
- Abaixo, há um **gerador de referência** para formatos comuns (Livro, Artigo, Site) — cole o resultado no bloco de referências do seu DOCX.
    """
)

st.subheader("Gerador rápido de referências (NBR 6023 – básico)")
ref_tipo = st.selectbox("Tipo", ["Livro", "Artigo", "Site"]) 
if ref_tipo == "Livro":
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor", "SILVA")
        ini = st.text_input("Iniciais", "J. P.")
        titulo = st.text_input("Título", "Métodos de pesquisa")
        ed = st.text_input("Edição (ex.: 2. ed.)", "")
    with c2:
        local = st.text_input("Local", "São Paulo")
        edit = st.text_input("Editora", "Atlas")
        ano = st.text_input("Ano", "2020")
    if st.button("Gerar referência (Livro)"):
        st.code(format_reference_livro(sb, ini, titulo, ed, local, edit, ano))
elif ref_tipo == "Artigo":
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor", "PEREIRA")
        ini = st.text_input("Iniciais", "M. A.")
        titulo = st.text_input("Título do artigo", "Sono e memória")
        period = st.text_input("Periódico", "Revista de Psicologia")
    with c2:
        vol = st.text_input("Volume", "37")
        num = st.text_input("Número (opcional)", "1")
        pags = st.text_input("Páginas", "63–76")
        ano = st.text_input("Ano", "2017")
    if st.button("Gerar referência (Artigo)"):
        st.code(format_reference_artigo(sb, ini, titulo, period, vol, num, pags, ano))
else:
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor (opcional)", "")
        ini = st.text_input("Iniciais (opcional)", "")
        titulo = st.text_input("Título da página", "Atendimento de Saúde Mental")
    with c2:
        site = st.text_input("Site", "Prefeitura de São José do Rio Preto")
        url = st.text_input("URL", "https://www.riopreto.sp.gov.br/cartaservicos/saude/atendimento-de-saude-mental-para-alcool-e-outras-drogas")
        acesso = st.text_input("Acesso em (DD MMM. AAAA)", "3 nov. 2025")
        ano = st.text_input("Ano (opcional)", "")
    if st.button("Gerar referência (Site)"):
        st.code(format_reference_site(sb if sb else None, ini if ini else None, titulo, site, url, acesso, ano if ano else None))

st.caption("💡 Dica: mantenha títulos como Heading 1/2/3 no Word; use marcadores para citações longas e referências; revise manualmente capas/sumários.")
"""
Streamlit ABNT Formatter (DOCX → DOCX)
--------------------------------------
- Aplica formatação ABNT em .docx e permite baixar o resultado.
- Foco: **NBR 14724** (margens, fonte, espaçamento, recuo, alinhamento, numeração) + **tabelas/figuras/citações longas** e **Referências (NBR 6023) – modo assistido**.

Como rodar (local):
1) pip install -r requirements.txt
2) streamlit run app.py

requirements.txt:
------------------
streamlit
python-docx
lxml

Notas:
- O app não exporta PDF (faça no Word/LibreOffice após download).
- "Citação longa" e "Referências NBR 6023" têm muita nuance: oferecemos **marcação assistida** e **templates**.
"""

import io
from typing import Optional, List, Tuple
import re
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================
# ABNT Helper Functions
# =====================

def set_page_margins(doc: Document, top_cm=3.0, left_cm=3.0, right_cm=2.0, bottom_cm=2.0):
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.bottom_margin = Cm(bottom_cm)


def configure_default_style(doc: Document, font_name="Times New Roman", font_size_pt=12, line_spacing=1.5,
                            first_line_indent_cm=1.25, justify=True):
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    pf = normal.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(first_line_indent_cm)


def style_all_paragraphs(doc: Document, justify=True, first_line_indent_cm=1.25):
    for p in doc.paragraphs:
        text = p.text.strip()
        # Headings: alinhar à esquerda, sem recuo
        if p.style and p.style.name.lower().startswith("heading"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            if justify:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def uppercase_heading_runs(paragraph):
    for run in paragraph.runs:
        run.text = run.text.upper()


def configure_heading_styles(doc: Document, h1_caps=True, h2_caps=True, h3_caps=True):
    heading_map = {"Heading 1": h1_caps, "Heading 2": h2_caps, "Heading 3": h3_caps}
    for p in doc.paragraphs:
        name = p.style.name if p.style is not None else ""
        if name in heading_map and heading_map[name]:
            uppercase_heading_runs(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)


def add_page_number_to_footer(doc: Document, position="right"):
    for section in doc.sections:
        footer = section.footer
        para = footer.add_paragraph() if len(footer.paragraphs) == 0 else footer.paragraphs[0]
        para.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT
        }.get(position, WD_ALIGN_PARAGRAPH.RIGHT)

        run = para.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = ' PAGE '
        fld_separate = OxmlElement('w:fldChar'); fld_separate.set(qn('w:fldCharType'), 'separate')
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr_text); run._r.append(fld_separate); run._r.append(fld_end)


def remove_extra_blank_lines(doc: Document):
    i = 0
    while i < len(doc.paragraphs) - 1:
        if not doc.paragraphs[i].text.strip() and not doc.paragraphs[i+1].text.strip():
            try:
                p = doc.paragraphs[i+1]._element
                p.getparent().remove(p)
                continue
            except Exception:
                pass
        i += 1

# ----------
# Tabelas
# ----------

def _set_row_cant_split(row):
    # <w:trPr><w:cantSplit/></w:trPr>
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)


def prevent_table_row_split_and_repeat_header(table):
    # Impede quebra de linha da linha na página seguinte e repete cabeçalho
    for i, row in enumerate(table.rows):
        _set_row_cant_split(row)
        if i == 0:
            trPr = row._tr.get_or_add_trPr()
            hdr = OxmlElement('w:tblHeader')
            trPr.append(hdr)
    # Melhor legibilidade de largura (evita autofit extremo)
    try:
        tblPr = table._tbl.get_or_add_tblPr()
        tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
    except Exception:
        pass


def center_paragraphs_with_drawings(doc: Document):
    # Centraliza parágrafos que contêm imagens (desenhos)
    for p in doc.paragraphs:
        if p._element.xpath('.//w:drawing'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)

# ----------
# Citações Longas (NBR 10520 – assistida)
# ----------
# Marcadores no texto: [[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]
# O app converterá os parágrafos marcados em bloco com: recuo 4 cm, fonte 10 pt, espaçamento simples, sem aspas.


def apply_long_quote_style(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(4)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.style.font.size = Pt(10) if paragraph.style and paragraph.style.font else None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def process_long_quote_markers(doc: Document) -> int:
    in_block = False
    changed = 0
    for p in doc.paragraphs:
        txt = p.text
        if '[[CITACAO_LONGA]]' in txt:
            in_block = True
            p.text = txt.replace('[[CITACAO_LONGA]]', '')
            apply_long_quote_style(p)
            changed += 1
            continue
        if '[[/CITACAO_LONGA]]' in txt:
            in_block = False
            p.text = txt.replace('[[/CITACAO_LONGA]]', '')
            apply_long_quote_style(p)
            changed += 1
            continue
        if in_block:
            apply_long_quote_style(p)
            changed += 1
    return changed

# ----------
# Referências (NBR 6023 – assistidas)
# ----------
# Opção 1: marcar bloco de referências com [[REFERENCIAS]] ... [[/REFERENCIAS]] para aplicar recuo francês
# Opção 2: utilizar gerador de referência por tipo (Livro, Artigo, Site)


def apply_references_block_format(doc: Document, first_line_hanging_cm=1.25, line_spacing=1.0, space_between_pts=6):
    in_refs = False
    count = 0
    for p in doc.paragraphs:
        t = p.text
        if '[[REFERENCIAS]]' in t:
            in_refs = True
            p.text = t.replace('[[REFERENCIAS]]', '')
        elif '[[/REFERENCIAS]]' in t:
            in_refs = False
            p.text = t.replace('[[/REFERENCIAS]]', '')
        if in_refs or '[[/REFERENCIAS]]' in t:
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(first_line_hanging_cm)
            pf.line_spacing = line_spacing
            pf.space_after = Pt(space_between_pts)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            count += 1
    return count


def format_reference_livro(autor_sobrenome: str, autor_iniciais: str, titulo: str, ed: Optional[str], local: str, editora: str, ano: str):
    # SOBRENOME, Iniciais. Título: subtítulo. ed. Local: Editora, ano.
    ed_str = f" {ed}." if ed else "."
    return f"{autor_sobrenome.upper()}, {autor_iniciais}. {titulo}.{ed_str} {local}: {editora}, {ano}."


def format_reference_artigo(autor_sobrenome: str, autor_iniciais: str, titulo: str, periodico: str, volume: str, numero: Optional[str], paginas: str, ano: str):
    num = f"({numero})" if numero else ""
    return f"{autor_sobrenome.upper()}, {autor_iniciais}. {titulo}. {periodico}, v. {volume} {num}, p. {paginas}, {ano}."


def format_reference_site(autor_sobrenome: Optional[str], autor_iniciais: Optional[str], titulo: str, site: str, url: str, acesso_data: str, ano: Optional[str]=None):
    autor = f"{autor_sobrenome.upper()}, {autor_iniciais}. " if (autor_sobrenome and autor_iniciais) else ""
    ano_str = f" {ano}." if ano else "."
    return f"{autor}{titulo}. {site}. Disponível em: <{url}>. Acesso em: {acesso_data}.{ano_str}"

# ----------
# Figuras (centralizar e legenda abaixo) e Tabelas (título acima, fonte abaixo)
# ----------


def add_caption_after_paragraph(p, text: str, italic=False):
    cap = p.insert_paragraph_after(text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    if italic and cap.runs:
        cap.runs[0].italic = True
    return cap


def ensure_captions(doc: Document, add_fig_captions: bool, add_tab_captions: bool):
    fig_n = 0
    tab_n = 0
    for block in doc.element.body:
        tag = block.tag
        if tag.endswith('}p'):
            p = next((pp for pp in doc.paragraphs if pp._p is block), None)
            if p is not None and p._element.xpath('.//w:drawing'):
                # imagem encontrada
                if add_fig_captions:
                    fig_n += 1
                    add_caption_after_paragraph(p, f"Figura {fig_n} – Descrição da figura", italic=False)
        elif tag.endswith('}tbl'):
            # tabela encontrada
            tbl = next((t for t in doc.tables if t._tbl is block), None)
            if tbl is not None:
                prevent_table_row_split_and_repeat_header(tbl)
                if add_tab_captions:
                    # Título acima e fonte abaixo (placeholder)
                    # Inserir título acima
                    first_p = doc.paragraphs[0]
                    new_p_above = first_p._element.__class__('w:p')
                    block.addprevious(new_p_above)
                    p_obj_above = None
                    for pp in doc.paragraphs:
                        if pp._p is new_p_above:
                            p_obj_above = pp; break
                    if p_obj_above:
                        p_obj_above.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_obj_above.paragraph_format.first_line_indent = Cm(0)
                        p_obj_above.add_run(f"Tabela {tab_n+1} – Título da tabela")
                    # Fonte abaixo
                    new_p_below = first_p._element.__class__('w:p')
                    block.addnext(new_p_below)
                    p_obj_below = None
                    for pp in doc.paragraphs:
                        if pp._p is new_p_below:
                            p_obj_below = pp; break
                    if p_obj_below:
                        p_obj_below.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_obj_below.paragraph_format.first_line_indent = Cm(0)
                        p_obj_below.add_run("Fonte: elaboração própria.")
                    tab_n += 1

# =====================
# Pipeline principal
# =====================

def apply_abnt_formatting(doc: Document,
                          h1_caps=True,
                          h2_caps=True,
                          h3_caps=True,
                          justify=True,
                          footer_page_numbers=True,
                          first_line_indent_cm=1.25,
                          center_images=True,
                          auto_captions_fig=True,
                          auto_captions_tab=True,
                          format_refs_block=True):
    set_page_margins(doc)
    configure_default_style(doc, line_spacing=1.5, first_line_indent_cm=first_line_indent_cm)
    style_all_paragraphs(doc, justify=justify, first_line_indent_cm=first_line_indent_cm)
    configure_heading_styles(doc, h1_caps=h1_caps, h2_caps=h2_caps, h3_caps=h3_caps)

    if center_images:
        center_paragraphs_with_drawings(doc)

    # Tabelas: impedir quebra de linha e repetir cabeçalho; adicionar legendas (opcional)
    for t in doc.tables:
        prevent_table_row_split_and_repeat_header(t)

    ensure_captions(doc, add_fig_captions=auto_captions_fig, add_tab_captions=auto_captions_tab)

    # Citações longas via marcadores
    process_long_quote_markers(doc)

    # Bloco de referências formatado
    if format_refs_block:
        apply_references_block_format(doc)

    if footer_page_numbers:
        add_page_number_to_footer(doc)

    remove_extra_blank_lines(doc)
    return doc

# =====================
# Streamlit UI
# =====================

st.set_page_config(page_title="ABNT Formatter (DOCX)", page_icon="📄", layout="centered")

st.title("📄 Automatizador ABNT para DOCX – Completo")
st.caption("Margens, fonte, parágrafos, títulos, numeração, **tabelas**, **figuras**, **citações longas** e **Referências (assistido)**.")

with st.expander("⚙️ Opções de formatação", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        do_h1 = st.checkbox("TÍTULOS NÍVEL 1 EM CAIXA ALTA", value=True)
        do_h2 = st.checkbox("TÍTULOS NÍVEL 2 EM CAIXA ALTA", value=True)
        do_h3 = st.checkbox("TÍTULOS NÍVEL 3 EM CAIXA ALTA", value=True)
        do_justify = st.checkbox("Justificar parágrafos", value=True)
        first_indent = st.number_input("Recuo da 1ª linha (cm)", min_value=0.0, max_value=3.0, value=1.25, step=0.25)
    with col2:
        do_pagenum = st.checkbox("Inserir número de página no rodapé (direita)", value=True)
        center_imgs = st.checkbox("Centralizar imagens (parágrafos com figuras)", value=True)
        auto_fig = st.checkbox("Adicionar legenda abaixo das figuras (placeholder)", value=False)
        auto_tab = st.checkbox("Adicionar título acima e fonte abaixo das tabelas (placeholder)", value=False)
        refs_block = st.checkbox("Formatar bloco de referências (recuo francês)", value=True)

st.markdown(
    """**Marcadores úteis no texto do DOCX** (opcionais):
- `[[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]` → aplica NBR 10520: recuo 4 cm, fonte ~10 pt, espaçamento simples, sem aspas.
- `[[REFERENCIAS]] ... [[/REFERENCIAS]]` → aplica recuo francês, espaçamento simples e espaço entre entradas.
    """
)

uploaded = st.file_uploader("Envie seu arquivo .docx", type=["docx"]) 

if uploaded is not None:
    try:
        doc = Document(uploaded)
    except Exception as e:
        st.error(f"Erro ao abrir DOCX: {e}")
        st.stop()

    st.info("Clique em **Aplicar ABNT** para processar. Você poderá baixar o DOCX formatado.")

    if st.button("✨ Aplicar ABNT"):
        try:
            formatted = apply_abnt_formatting(
                doc,
                h1_caps=do_h1,
                h2_caps=do_h2,
                h3_caps=do_h3,
                justify=do_justify,
                footer_page_numbers=do_pagenum,
                first_line_indent_cm=float(first_indent),
                center_images=center_imgs,
                auto_captions_fig=auto_fig,
                auto_captions_tab=auto_tab,
                format_refs_block=refs_block,
            )
            out = io.BytesIO(); formatted.save(out); out.seek(0)
            base_name = uploaded.name.replace('.docx', '').replace('.DOCX', '')
            st.success("Arquivo formatado com sucesso!")
            st.download_button(
                label="⬇️ Baixar DOCX formatado",
                data=out,
                file_name=f"{base_name}_ABNT.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.caption("Dica: revise capa/folha de rosto, listas pré-textuais e sumário no Word/LibreOffice e exporte para PDF.")
        except Exception as e:
            st.error(f"Falha ao formatar: {e}")

st.divider()

st.subheader("Cobertura desta automação")
st.markdown(
    """
**Layout geral (NBR 14724):**
- **Margens**: 3 cm (superior/esquerda) e 2 cm (direita/inferior)
- **Fonte**: Times New Roman, 12 pt (texto)
- **Parágrafos**: 1,5; sem espaços extras; **recuo 1,25 cm**; justificado (opcional)
- **Títulos (Heading 1/2/3)**: caixa alta opcional; esquerda; sem recuo
- **Páginas**: numeração no rodapé (direita) com campo automático `PAGE`

**Tabelas e Figuras:**
- **Tabelas**: impede que a **linha quebre entre páginas** e repete **cabeçalho**; opção de **Título acima** e **Fonte abaixo** (placeholders).
- **Figuras**: centraliza parágrafo com imagem; opção de **Legenda abaixo** numerada (placeholder).

**Citações (NBR 10520):**
- **Citação curta**: siga no corpo entre aspas (até ~3 linhas) – revisão manual.
- **Citação longa**: use os marcadores `[[CITACAO_LONGA]] ... [[/CITACAO_LONGA]]` para aplicar bloco ABNT.

**Referências (NBR 6023):**
- Use `[[REFERENCIAS]] ... [[/REFERENCIAS]]` para aplicar **recuo francês**, espaçamento simples e espaço entre entradas.
- Abaixo, há um **gerador de referência** para formatos comuns (Livro, Artigo, Site) — cole o resultado no bloco de referências do seu DOCX.
    """
)

st.subheader("Gerador rápido de referências (NBR 6023 – básico)")
ref_tipo = st.selectbox("Tipo", ["Livro", "Artigo", "Site"]) 
if ref_tipo == "Livro":
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor", "SILVA")
        ini = st.text_input("Iniciais", "J. P.")
        titulo = st.text_input("Título", "Métodos de pesquisa")
        ed = st.text_input("Edição (ex.: 2. ed.)", "")
    with c2:
        local = st.text_input("Local", "São Paulo")
        edit = st.text_input("Editora", "Atlas")
        ano = st.text_input("Ano", "2020")
    if st.button("Gerar referência (Livro)"):
        st.code(format_reference_livro(sb, ini, titulo, ed, local, edit, ano))
elif ref_tipo == "Artigo":
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor", "PEREIRA")
        ini = st.text_input("Iniciais", "M. A.")
        titulo = st.text_input("Título do artigo", "Sono e memória")
        period = st.text_input("Periódico", "Revista de Psicologia")
    with c2:
        vol = st.text_input("Volume", "37")
        num = st.text_input("Número (opcional)", "1")
        pags = st.text_input("Páginas", "63–76")
        ano = st.text_input("Ano", "2017")
    if st.button("Gerar referência (Artigo)"):
        st.code(format_reference_artigo(sb, ini, titulo, period, vol, num, pags, ano))
else:
    c1, c2 = st.columns(2)
    with c1:
        sb = st.text_input("SOBRENOME do autor (opcional)", "")
        ini = st.text_input("Iniciais (opcional)", "")
        titulo = st.text_input("Título da página", "Atendimento de Saúde Mental")
    with c2:
        site = st.text_input("Site", "Prefeitura de São José do Rio Preto")
        url = st.text_input("URL", "https://www.riopreto.sp.gov.br/cartaservicos/saude/atendimento-de-saude-mental-para-alcool-e-outras-drogas")
        acesso = st.text_input("Acesso em (DD MMM. AAAA)", "3 nov. 2025")
        ano = st.text_input("Ano (opcional)", "")
    if st.button("Gerar referência (Site)"):
        st.code(format_reference_site(sb if sb else None, ini if ini else None, titulo, site, url, acesso, ano if ano else None))

st.caption("💡 Dica: mantenha títulos como Heading 1/2/3 no Word; use marcadores para citações longas e referências; revise manualmente capas/sumários.")
